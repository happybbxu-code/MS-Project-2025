import json
import sys
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(spec_path, out_path):
    with open(spec_path, 'r') as f:
        spec = json.load(f)
    
    doc = Document()
    
    # Basic page setup
    if 'page' in spec:
        p = spec['page']
        section = doc.sections[0]
        section.left_margin = Mm(p.get('left', 25.4))
        section.right_margin = Mm(p.get('right', 25.4))
        section.top_margin = Mm(p.get('top', 25.4))
        section.bottom_margin = Mm(p.get('bottom', 25.4))

    for block in spec['blocks']:
        btype = block['type']
        if btype == 'heading':
            doc.add_heading(block['text'], level=block['level'])
        elif btype == 'paragraph':
            doc.add_paragraph(block['text'])
        elif btype == 'bullet_list':
            for item in block['items']:
                doc.add_paragraph(item, style='List Bullet')
    
    doc.save(out_path)

if __name__ == "__main__":
    create_docx(sys.argv[1], sys.argv[2])
